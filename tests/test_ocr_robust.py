"""OCR 健壮性优化的纯函数回归：文本层乱码闸门 + 中文兜底择优的度量。"""
import unittest

from extraction.extract import pdf_type
from extraction.extract import ocr


class TextLayerSuspectTest(unittest.TestCase):
    def test_normal_text_not_suspect(self):
        self.assertFalse(pdf_type.text_layer_suspect(
            "Invoice No INV-1  Total 100.00  结算币种 USD  周啟邦律師事務所"))

    def test_garbled_cid_suspect(self):
        garble = "".join(chr(c) for c in range(0xE000, 0xE030))   # 私用区（未映射CID典型）
        self.assertTrue(pdf_type.text_layer_suspect(garble * 3))

    def test_replacement_char_suspect(self):
        self.assertTrue(pdf_type.text_layer_suspect("�" * 40))

    def test_short_text_not_suspect(self):
        self.assertFalse(pdf_type.text_layer_suspect(""))   # 太短不判


class UsefulCharsTest(unittest.TestCase):
    def test_cjk_weighted_over_latin(self):
        # 中文兜底择优：等量字符下，含中文的一方"有效字符"更高 → 胜出
        self.assertGreater(ocr._useful_chars("发票总额柒仟"), ocr._useful_chars("garbage"))

    def test_cjk_count(self):
        self.assertEqual(ocr._cjk_count("周啟邦 K.B. Chau 律師"), 5)
        self.assertEqual(ocr._cjk_count("English only 123"), 0)


class RegionTextTest(unittest.TestCase):
    def test_tiny_box_returns_empty_without_ocr(self):
        # 点击级的极小框：在调用 OCR 前就返回空（不误触发引擎）
        from PIL import Image
        import gateway.main as M
        img = Image.new("RGB", (800, 600), "white")
        self.assertEqual(M._ocr_region_text(img, 0.5, 0.5, 0.501, 0.501), "")


class LargeImageGuardTest(unittest.TestCase):
    def test_bound_large_downscales(self):
        from PIL import Image
        from extraction.extract import ocr
        out, f = ocr._bound_large(Image.new("RGB", (6000, 4000), "white"), max_side=4500)
        self.assertEqual(max(out.size), 4500)
        self.assertLess(f, 1.0)

    def test_bound_large_noop_when_small(self):
        from PIL import Image
        from extraction.extract import ocr
        out, f = ocr._bound_large(Image.new("RGB", (800, 600), "white"), max_side=4500)
        self.assertEqual(out.size, (800, 600))
        self.assertEqual(f, 1.0)

    def test_pil_pixel_cap_set(self):
        from PIL import Image
        from core import config
        self.assertLessEqual(Image.MAX_IMAGE_PIXELS * 2, config.MAX_IMAGE_PIXELS + 1)


class DosGuardTest(unittest.TestCase):
    def _pdf(self, pages):
        import tempfile, os, fitz
        p = os.path.join(tempfile.mkdtemp(), f"p{pages}.pdf")
        doc = fitz.open()
        for _ in range(pages):
            doc.new_page()
        doc.save(p); doc.close()
        import pathlib
        return pathlib.Path(p)

    def test_pdf_pages_over_limit_rejected(self):
        from extraction import pipeline
        from core import config
        with self.assertRaises(ValueError):
            pipeline._guard_pdf_pages(self._pdf(config.MAX_PDF_PAGES + 3))

    def test_pdf_pages_within_limit_ok(self):
        from extraction import pipeline
        pipeline._guard_pdf_pages(self._pdf(3))   # 不抛

    def test_embedded_images_cap_config(self):
        from core import config
        self.assertGreater(config.MAX_EMBEDDED_IMAGES, 0)   # 上限存在（切片逻辑在 _process_embedded_images）


class MultiPagePreviewTest(unittest.TestCase):
    def test_fill_all_page_sizes_preserves_first(self):
        import tempfile, os
        import fitz
        from core.models import Invoice
        from extraction import pipeline
        d = tempfile.mkdtemp()
        p = os.path.join(d, "two.pdf")
        doc = fitz.open()
        doc.new_page(width=300, height=400)
        doc.new_page(width=300, height=400)
        doc.save(p); doc.close()
        inv = Invoice(file_hash="h", file_name="two.pdf")
        inv.page_sizes = [[288.0, 384.0]]                 # 首页来自 OCR（略不同尺度）
        pipeline._fill_all_page_sizes(inv, __import__("pathlib").Path(p))
        self.assertEqual(len(inv.page_sizes), 2)          # 两页都填上
        self.assertEqual(inv.page_sizes[0], [288.0, 384.0])   # 首页保留 OCR 尺度


class OfficeImageFormTest(unittest.TestCase):
    """图嵌入带大段文字的 docx：真发票在图里，应判 image_form 走 OCR，不被说明文字带偏。"""
    def _docx(self, paras, with_image=True):
        import tempfile, os
        from docx import Document
        from docx.shared import Inches
        d = tempfile.mkdtemp(); p = os.path.join(d, "t.docx")
        doc = Document()
        for t in paras:
            doc.add_paragraph(t)
        if with_image:
            from PIL import Image
            ip = os.path.join(d, "i.png"); Image.new("RGB", (700, 400), "white").save(ip)
            doc.add_picture(ip, width=Inches(5))
        doc.save(p)
        import pathlib
        return pathlib.Path(p)

    def test_prose_plus_image_is_image_form(self):
        from extraction.extract import word
        p = self._docx(["This document is provided for your records. " * 4])
        self.assertTrue(word.is_image_form(p))          # 纯说明+图 → 走 OCR

    def test_text_invoice_not_image_form(self):
        from extraction.extract import word
        p = self._docx(["Invoice No INV-100  Total Due USD 500.00  Date 2026-05-01 " * 2])
        self.assertFalse(word.is_image_form(p))         # 有发票要素 → 文本路径（不回归）

    def test_no_image_not_image_form(self):
        from extraction.extract import word
        p = self._docx(["short", ], with_image=False)
        self.assertFalse(word.is_image_form(p))         # 无图 → 不是图片型



class RegionWordsFastPathTest(unittest.TestCase):
    def test_pick_words_by_box(self):
        import gateway.main as M
        words = [[0, 0.1, 0.1, 0.3, 0.15, "Invoice"], [0, 0.35, 0.1, 0.6, 0.15, "INV-9"],
                 [0, 0.1, 0.8, 0.4, 0.85, "Total"], [1, 0.1, 0.1, 0.3, 0.15, "page2"]]
        # 顶部行两个词都取到，按 x 排序拼接
        self.assertEqual(M._pick_words_text(words, 0, 0.0, 0.05, 0.7, 0.2), "Invoice INV-9")
        # 底部只取 Total
        self.assertEqual(M._pick_words_text(words, 0, 0.0, 0.75, 0.5, 0.9), "Total")
        # 第 0 页不含第 1 页的词
        self.assertNotIn("page2", M._pick_words_text(words, 0, 0.0, 0.0, 1.0, 1.0))
        # 空区域
        self.assertEqual(M._pick_words_text(words, 0, 0.9, 0.9, 0.95, 0.95), "")



class AmountTokenFixTest(unittest.TestCase):
    def _box(self, x0, x1):
        return [[x0, 10], [x1, 10], [x1, 30], [x0, 30]]

    def test_standalone_S_with_adjacent_money_to_dollar(self):
        from extraction.extract import ocr
        toks = [("S", 0.8, self._box(10, 25)), ("7,000.00", 0.9, self._box(30, 90))]
        out = {t for t, _, _ in ocr._fix_amount_tokens(toks)}
        self.assertIn("$", out)

    def test_S_glued_money_to_dollar(self):
        from extraction.extract import ocr
        out = {t for t, _, _ in ocr._fix_amount_tokens([("S7,000.00", 0.9, self._box(10, 90))])}
        self.assertIn("$7,000.00", out)

    def test_colon_in_money_to_dot(self):
        from extraction.extract import ocr
        out = {t for t, _, _ in ocr._fix_amount_tokens([("$7,000:00", 0.9, self._box(10, 90))])}
        self.assertIn("$7,000.00", out)

    def test_plain_S_not_touched(self):
        from extraction.extract import ocr  # 无邻近金额的 S（如 Solicitors 附近）不误纠
        toks = [("S", 0.8, self._box(10, 25)), ("Solicitors", 0.9, self._box(400, 520))]
        self.assertEqual([t for t, _, _ in ocr._fix_amount_tokens(toks)], ["S", "Solicitors"])



class FromToLabelTest(unittest.TestCase):
    def test_from_to_labels(self):
        from extraction.parse import generic as g
        self.assertEqual(g._label_match("From")[0], "issuer_name")     # 整格 From
        self.assertEqual(g._label_match("To:")[0], "customer_name")    # To: 带冒号
        # 散文里的 from 不误命中开票方（不是整格 "From"）
        self.assertIsNone(g._label_match("payable from the bank account below"))



class BankSwiftLabelTest(unittest.TestCase):
    def test_swift_code_label_consumed(self):
        from extraction.parse import generic as g
        rows = [
            [(0.0, 60.0, "Bank Details")],
            [(0.0, 120.0, "SWIFT Code: HASEHKHH")],
        ]
        out = g.extract_bank(rows)
        self.assertEqual(out.get("bank_swift"), "HASEHKHH")   # "Code:" 不应漏进值



class IssuerNotLineItemTest(unittest.TestCase):
    def test_line_item_not_picked_as_issuer(self):
        from extraction.parse import generic as g
        # 无开票方公司名、只有明细行 → 不得把明细当开票方
        rows = [[(0.0, 60.0, "TAX INVOICE")], [(0.0, 90.0, "Consulting fee 100.00")]]
        name, _addr, _e, _p = g.extract_issuer(rows)
        self.assertNotEqual(name, "Consulting fee 100.00")


if __name__ == "__main__":
    unittest.main()


class TestNegativeSummaryCorrection(unittest.TestCase):
    """非贷记单汇总不应为负 → 小计/税/合计取正（OCR 印章弧线读成负号，或文本件括号重述被当会计负数）。"""

    def _inv(self, sub, ocr=True):
        from core.models import Invoice, FieldValue
        from decimal import Decimal
        inv = Invoice(ocr_used=ocr)
        inv.set("subtotal", FieldValue(raw="-3,562.00", value=Decimal(sub)))
        inv.set("total_due", FieldValue(raw="4,274.40", value=Decimal("4274.40")))
        return inv

    def test_ocr_negative_subtotal_flipped_positive(self):
        from extraction.parse.template_rules import _correct_negative_summaries
        from decimal import Decimal
        inv = self._inv("-3562.00")
        _correct_negative_summaries(inv, is_credit=False)
        self.assertEqual(inv.f("subtotal").value, Decimal("3562.00"))
        self.assertEqual(inv.f("subtotal").raw, "3,562.00")   # raw 也去掉前导负号

    def test_credit_note_keeps_negative(self):
        from extraction.parse.template_rules import _correct_negative_summaries
        from decimal import Decimal
        inv = self._inv("-3562.00")
        _correct_negative_summaries(inv, is_credit=True)      # 贷记单：保留负数
        self.assertEqual(inv.f("subtotal").value, Decimal("-3562.00"))

    def test_text_pdf_also_corrected(self):
        from extraction.parse.template_rules import _correct_negative_summaries
        from decimal import Decimal
        # 文本件（非 OCR）也取正：如 "Total Due: ... ($2,500.00)" 括号被当会计负数
        inv = self._inv("-3562.00", ocr=False)
        _correct_negative_summaries(inv, is_credit=False)
        self.assertEqual(inv.f("subtotal").value, Decimal("3562.00"))


class TestRecoverObscuredTotals(unittest.TestCase):
    """水印/印章糊掉 Total/Tax 标签、数值作为孤立行幸存 → 按算术一致性归位（仅 OCR）。"""

    def _inv(self, sub, tax=None, ocr=True):
        from core.models import Invoice, FieldValue
        from decimal import Decimal
        inv = Invoice(ocr_used=ocr)
        inv.set("subtotal", FieldValue(raw=sub, value=Decimal(sub.replace(",", ""))))
        if tax is not None:
            inv.set("sales_tax", FieldValue(raw=tax, value=Decimal(tax.replace(",", ""))))
        return inv

    def test_single_orphan_equals_subtotal_is_total(self):
        from extraction.parse.template_rules import _recover_obscured_totals
        from decimal import Decimal
        inv = self._inv("4225.00")
        full = "Subtotal $4,225.00\nVOIDVOD\n$4,225.00\n"     # 合计标签被糊、值幸存
        _recover_obscured_totals(inv, full)
        self.assertEqual(inv.f("total_due").value, Decimal("4225.00"))
        self.assertIn("遮盖", inv.f("total_due").note)

    def test_two_orphans_recover_tax_and_total(self):
        from extraction.parse.template_rules import _recover_obscured_totals
        from decimal import Decimal
        inv = self._inv("5346.00")                            # 税+合计标签都被糊
        full = "Subtotal GBP 5,346.00\nGBP 374.22\nGBP 5,720.22\n"
        _recover_obscured_totals(inv, full)
        self.assertEqual(inv.f("sales_tax").value, Decimal("374.22"))
        self.assertEqual(inv.f("total_due").value, Decimal("5720.22"))

    def test_tax_known_pick_matching_total(self):
        from extraction.parse.template_rules import _recover_obscured_totals
        from decimal import Decimal
        inv = self._inv("5346.00", tax="374.22")
        full = "Subtotal GBP 5,346.00\nTax GBP 374.22\nVOID\nGBP 5,720.22\n"
        _recover_obscured_totals(inv, full)
        self.assertEqual(inv.f("total_due").value, Decimal("5720.22"))

    def test_no_arithmetic_match_leaves_missing(self):
        from extraction.parse.template_rules import _recover_obscured_totals
        inv = self._inv("4225.00")
        full = "Subtotal $4,225.00\nVOID\n$9,999.99\n"        # 孤立值与小计不自洽 → 不赋值
        _recover_obscured_totals(inv, full)
        self.assertIsNone(inv.f("total_due").value)

    def test_text_pdf_skipped(self):
        from extraction.parse.template_rules import _recover_obscured_totals
        inv = self._inv("4225.00", ocr=False)                 # 文本 PDF 不做恢复
        _recover_obscured_totals(inv, "Subtotal $4,225.00\n$4,225.00\n")
        self.assertIsNone(inv.f("total_due").value)

    def test_existing_total_not_overwritten(self):
        from extraction.parse.template_rules import _recover_obscured_totals
        from core.models import FieldValue
        from decimal import Decimal
        inv = self._inv("4225.00")
        inv.set("total_due", FieldValue(raw="4,225.00", value=Decimal("4225.00")))
        _recover_obscured_totals(inv, "Subtotal $4,225.00\n$9,999.99\n")   # 已有值不动
        self.assertEqual(inv.f("total_due").value, Decimal("4225.00"))
