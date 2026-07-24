"""Excel(.xlsx) 发票解析：openpyxl 结构化读取 → 复用通用提取（日期/金额/明细/勾稽）。"""
import datetime
import shutil
import tempfile
import unittest
from decimal import Decimal
from pathlib import Path

from core import config, db


def _has_openpyxl():
    try:
        import openpyxl  # noqa
        return True
    except Exception:
        return False


@unittest.skipUnless(_has_openpyxl(), "需要 openpyxl")
class ExcelInvoiceTest(unittest.TestCase):
    def setUp(self):
        self._dir = tempfile.mkdtemp()
        self._db, self._up = config.DB_PATH, config.UPLOAD_DIR
        config.DB_PATH = Path(self._dir) / "t.db"
        config.UPLOAD_DIR = Path(self._dir) / "up"
        config.UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
        db._initialized = False
        db.init_db()

    def tearDown(self):
        config.DB_PATH, config.UPLOAD_DIR = self._db, self._up
        db._initialized = False
        shutil.rmtree(self._dir, ignore_errors=True)

    def _make_table_xlsx(self) -> Path:
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws["B2"] = "Acme Spreadsheet Labs Inc."
        ws["B4"] = "Invoice No."; ws["C4"] = "XL-2026-001"
        ws["B5"] = "Invoice Date"
        ws["C5"] = datetime.datetime(2026, 7, 3); ws["C5"].number_format = "yyyy-mm-dd"
        ws["B7"] = "Bill To"
        ws["B8"] = "Metro Holdings"
        ws["B10"] = "DESCRIPTION"; ws["C10"] = "QTY"; ws["D10"] = "UNIT PRICE"; ws["E10"] = "AMOUNT"
        rows = [("Implementation workshop", 1, 925, 925),
                ("Data reconciliation support", 5, 168, 840)]
        for i, (d, q, u, a) in enumerate(rows):
            r = 11 + i
            ws[f"B{r}"] = d; ws[f"C{r}"] = q
            ws[f"D{r}"] = u; ws[f"D{r}"].number_format = "$#,##0.00"
            ws[f"E{r}"] = a; ws[f"E{r}"].number_format = "$#,##0.00"
        ws["D14"] = "Subtotal"; ws["E14"] = 1765; ws["E14"].number_format = "$#,##0.00"
        ws["D15"] = "Total Due"; ws["E15"] = 1765; ws["E15"].number_format = "$#,##0.00"
        p = Path(self._dir) / "table.xlsx"
        wb.save(p)
        return p

    def test_table_excel_extraction_and_reconciliation(self):
        from extraction import pipeline
        inv = pipeline.process_local(self._make_table_xlsx())[0]
        self.assertEqual(inv.parse_method, "excel")
        self.assertEqual(inv.f("invoice_no").value, "XL-2026-001")
        self.assertEqual(inv.f("invoice_date").value, "2026-07-03")     # 日期非序列号
        self.assertEqual(len(inv.line_items), 2)
        self.assertEqual(inv.line_items[0].description, "Implementation workshop")
        self.assertEqual(inv.line_items[0].quantity, Decimal("1"))      # 数量未并入描述
        self.assertEqual(inv.line_items[0].amount, Decimal("925"))
        # 明细合计 == 小计（勾稽）
        self.assertEqual(sum(li.amount for li in inv.line_items), Decimal("1765"))
        # 字段有 bbox（与自渲染 PNG 对齐 → 可双向高亮）
        self.assertIsNotNone(inv.f("invoice_no").bbox)
        b = inv.f("invoice_no").bbox
        pw, ph = inv.page_sizes[0]
        self.assertTrue(0 <= b[1] <= pw and 0 <= b[2] <= ph)   # 框在页内

    def test_render_uses_bundled_truetype_font_size_honored(self):
        """Excel 自渲染必须用随包 TrueType 字体（跨平台一致）、字号生效——
        不能退回 PIL 位图默认字体（那会忽略字号，字体大小全乱：Mac 上"该大的小该小的大"）。"""
        from extraction.extract import excel
        f_small, f_big = excel._font(False, 10), excel._font(False, 30)
        self.assertEqual(type(f_small).__name__, "FreeTypeFont")     # 非 load_default 位图字体
        self.assertGreater(f_big.getbbox("Invoice")[2], f_small.getbbox("Invoice")[2] * 2)
        # 随包字体文件存在（打包会带上）
        from pathlib import Path as _P
        self.assertTrue((_P(excel.__file__).parent / "fonts" / "DejaVuSans.ttf").exists())

    def test_excel_render_png_matches_page_size(self):
        from extraction.extract import excel
        import io
        from PIL import Image
        p = self._make_table_xlsx()
        doc = excel.excel_to_pdfdoc(p)
        pw, ph = doc.page_sizes[0]
        im = Image.open(io.BytesIO(excel.render_png(p, scale=2.0)))
        self.assertAlmostEqual(im.size[0], pw * 2.0, delta=2)   # PNG 与 page_size 同坐标系
        self.assertAlmostEqual(im.size[1], ph * 2.0, delta=2)

    def test_excel_in_upload_whitelist(self):
        self.assertIn(".xlsx", config.ALLOWED_UPLOAD_EXTS)

    def _mini_invoice(self, ws, no, total, start_row=1):
        """在 ws 从 start_row 起写一张极简发票（标签+值相邻），返回写到的末行。"""
        import datetime
        rows = [("Invoice No.", no), ("Invoice Date", datetime.datetime(2026, 7, 1)),
                ("Total Due", total)]
        for i, (lab, val) in enumerate(rows):
            r = start_row + i
            ws.cell(row=r, column=2, value=lab)
            c = ws.cell(row=r, column=3, value=val)
            if lab == "Total Due":
                c.number_format = "$#,##0.00"
        return start_row + len(rows) - 1

    def test_multi_sheet_excel_splits_into_invoices(self):
        """一个 xlsx 多个工作表、每表一张发票 → 拆成多张独立记录。"""
        import openpyxl
        from extraction import pipeline
        wb = openpyxl.Workbook()
        ws1 = wb.active; ws1.title = "INV-A"; self._mini_invoice(ws1, "INV-1001", 100)
        ws2 = wb.create_sheet("INV-B"); self._mini_invoice(ws2, "INV-1002", 200)
        p = Path(self._dir) / "multi_sheet.xlsx"; wb.save(p)
        invs = pipeline.process_local(p)
        self.assertEqual(len(invs), 2)
        nos = sorted(i.f("invoice_no").value for i in invs)
        self.assertEqual(nos, ["INV-1001", "INV-1002"])

    def test_stacked_single_sheet_excel_splits(self):
        """单个工作表内纵向堆叠多张发票（多个 TOTAL + Invoice No 锚点）→ 按段拆分。"""
        import openpyxl
        from extraction import pipeline
        wb = openpyxl.Workbook(); ws = wb.active; ws.title = "Both"
        end = self._mini_invoice(ws, "INV-2001", 100, start_row=1)
        self._mini_invoice(ws, "INV-2002", 200, start_row=end + 3)   # 隔空再来一张
        p = Path(self._dir) / "stacked.xlsx"; wb.save(p)
        invs = pipeline.process_local(p)
        self.assertEqual(len(invs), 2)
        self.assertEqual(sorted(i.f("invoice_no").value for i in invs), ["INV-2001", "INV-2002"])

    def test_single_excel_not_split(self):
        """单张发票的 xlsx 不被误拆（invoice_units 返回 None）。"""
        from extraction.extract import excel
        p = self._make_table_xlsx()
        self.assertIsNone(excel.invoice_units(p))

    @unittest.skipUnless(__import__("importlib").util.find_spec("docx"), "需要 python-docx")
    def test_multi_invoice_word_physical_split(self):
        """多发票 Word（即便被 fitz 折叠到一页）→ python-docx 物理拆成一文件一发票，
        每张走单张 Word 路径（可渲染原件，与 PDF/Excel 同思路）。"""
        try:
            import fitz  # noqa
        except Exception:
            self.skipTest("需要 PyMuPDF")
        import docx
        from extraction import pipeline
        d = docx.Document()
        for no, tot in [("WD-1001", "100.00"), ("WD-2002", "200.00")]:
            d.add_paragraph("INVOICE")
            d.add_paragraph("Northbridge Labs LLC")
            d.add_paragraph("1400 Market Test Ave, Suite 210, Seattle, WA 98101")
            d.add_paragraph(f"Invoice No.: {no}")
            d.add_paragraph("Invoice Date: 2026-07-01")
            d.add_paragraph("Service: consulting and advisory services for the billing period")
            d.add_paragraph(f"Total Due: USD {tot}")
            d.add_page_break()
        p = Path(self._dir) / "multi.docx"; d.save(p)
        invs = pipeline.process_local(p)
        self.assertEqual(len(invs), 2)
        nos = [i.f("invoice_no").value for i in invs]
        self.assertIn("WD-1001", nos); self.assertIn("WD-2002", nos)
        # 物理拆分：每张是独立可渲染 docx（有 page_sizes），而非"归档文本"降级
        self.assertTrue(all(i.parse_method == "pdf_text" and i.page_sizes for i in invs))
        self.assertTrue(all(i.file_name.endswith(".docx") for i in invs))

    def test_single_word_not_split(self):
        """单张 Word 不被误拆（invoice_units 返回 None）。"""
        if __import__("importlib").util.find_spec("docx") is None:
            self.skipTest("需要 python-docx")
        import docx
        from extraction.extract import word
        d = docx.Document()
        d.add_paragraph("INVOICE"); d.add_paragraph("Invoice No.: SOLO-1")
        d.add_paragraph("Total Due: USD 500.00")
        p = Path(self._dir) / "solo.docx"; d.save(p)
        self.assertIsNone(word.invoice_units(p))


    def _png(self, size, color="white"):
        from PIL import Image
        p = Path(self._dir) / f"img_{size[0]}x{size[1]}_{color}.png"
        Image.new("RGB", size, color).save(p)
        return p

    def test_excel_extract_images_filters_logo(self):
        """xlsx 提取内嵌位图：保留大发票图、过滤小 logo。"""
        import openpyxl
        from openpyxl.drawing.image import Image as XLImage
        from extraction.extract import excel
        wb = openpyxl.Workbook(); ws = wb.active
        ws.add_image(XLImage(str(self._png((800, 1000)))), "A1")    # 发票大图
        ws.add_image(XLImage(str(self._png((60, 60), "blue"))), "H1")  # logo
        p = Path(self._dir) / "imgform.xlsx"; wb.save(p)
        imgs = excel.extract_images(p)
        self.assertEqual(len(imgs), 1)                              # 只保留大图
        self.assertTrue(excel.is_image_form(p))                    # 几乎无文本+有图 → 图片形式

    def test_word_extract_images_filters_logo(self):
        if __import__("importlib").util.find_spec("docx") is None:
            self.skipTest("需要 python-docx")
        import docx
        from extraction.extract import word
        d = docx.Document()
        d.add_picture(str(self._png((800, 1000))))
        d.add_picture(str(self._png((50, 50), "red")))
        p = Path(self._dir) / "imgform.docx"; d.save(p)
        self.assertEqual(len(word.extract_images(p)), 1)
        self.assertTrue(word.is_image_form(p))

    def test_text_excel_not_image_form(self):
        """正常文本/表格 xlsx 不被当作图片形式。"""
        from extraction.extract import excel
        self.assertFalse(excel.is_image_form(self._make_table_xlsx()))

    def test_text_excel_with_logo_still_text_path(self):
        """有发票文本 + 内嵌 logo/大图 → 仍走文本提取（不因有图就走图片路径）。"""
        import openpyxl
        from openpyxl.drawing.image import Image as XLImage
        from extraction.extract import excel
        from extraction import pipeline
        p = self._make_table_xlsx()                 # 已有完整发票文本
        wb = openpyxl.load_workbook(p)
        wb.active.add_image(XLImage(str(self._png((300, 120), "navy"))), "G2")   # logo
        wb.active.add_image(XLImage(str(self._png((800, 600), "teal"))), "G10")  # 甚至一张大图
        wb.save(p)
        self.assertFalse(excel.is_image_form(p))    # 有文本 → 不是图片形式
        inv = pipeline.process_local(p)[0]
        self.assertEqual(inv.parse_method, "excel") # 仍走 Excel 文本路径
        self.assertEqual(inv.f("invoice_no").value, "XL-2026-001")

    def test_text_word_with_logo_still_text_path(self):
        """有发票文本的 Word + 内嵌 logo → 仍走文本路径，不被当图片形式。"""
        if __import__("importlib").util.find_spec("docx") is None:
            self.skipTest("需要 python-docx")
        import docx
        from extraction.extract import word
        d = docx.Document()
        d.add_paragraph("INVOICE  Northbridge Labs LLC")
        d.add_paragraph("1400 Market Test Ave, Suite 210, Seattle, WA 98101")
        d.add_paragraph("Invoice No.: TXT-1   Invoice Date: 2026-07-01")
        d.add_paragraph("Consulting and advisory services for the billing period")
        d.add_paragraph("Total Due: USD 1,200.00")
        d.add_picture(str(self._png((300, 120), "navy")))      # logo
        p = Path(self._dir) / "text_logo.docx"; d.save(p)
        self.assertFalse(word.is_image_form(p))

    def test_image_form_excel_routes_each_image_to_image_path(self):
        """图片形式 xlsx 多张图 → 每张另存为图片、各走图片路径成独立记录。"""
        import openpyxl
        from openpyxl.drawing.image import Image as XLImage
        from extraction import pipeline
        wb = openpyxl.Workbook(); ws = wb.active
        ws.add_image(XLImage(str(self._png((800, 1000)))), "A1")
        ws.add_image(XLImage(str(self._png((800, 1000), "ivory"))), "A60")
        p = Path(self._dir) / "two_img.xlsx"; wb.save(p)
        invs = pipeline.process_local(p)
        self.assertEqual(len(invs), 2)                             # 一张图=一张发票
        self.assertTrue(all(i.file_name.lower().endswith(".png") for i in invs))  # 走图片路径


if __name__ == "__main__":
    unittest.main()
