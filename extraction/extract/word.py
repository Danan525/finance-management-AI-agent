"""Word(.docx) 多发票物理拆分：与 PDF/Excel 同思路——先拆成"一文件一发票"的 docx，
再各自走单张 Word 路径（fitz 渲染原件 + 字段提取 + bbox + 勾稽），体验一致。

为什么不靠 fitz 按页拆：docx 分页是动态的，短/连续的多发票常被 fitz 折叠到一页、拆不开。
故改用 python-docx 按**文档块（段落/表格）顺序 + 发票锚点**切分，物理生成多个单发票 docx
（deepcopy 块的 XML，保留加粗/颜色/表格等格式）。切不清楚（完整性校验不过）就不拆。
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import List, Optional, Tuple

_WORD_EXTS = (".docx", ".docm")
_TITLE_RE = re.compile(r"^\s*(tax\s+)?invoice\s*$", re.IGNORECASE)
_NO_RE = re.compile(r"invoice\s*(no|number|#)", re.IGNORECASE)
_MIN_IMG_AREA = 200_000          # 像素面积下限：≈ 500×400 以上才算"整页发票图"，滤掉 logo/banner/装饰图
_MIN_TEXT_CHARS = 100            # 可提取文本少于此 → 视为"图片形式发票"
_RASTER = {"PNG": ".png", "JPEG": ".jpg", "BMP": ".bmp", "TIFF": ".tif", "GIF": ".gif"}


def is_word(path) -> bool:
    return Path(path).suffix.lower() in _WORD_EXTS


def extract_images(path) -> List[Tuple[bytes, str]]:
    """提取 docx 内嵌的**位图发票图**（过滤过小的 logo/装饰、跳过矢量图）→ [(bytes, ext), ...]。

    缺 python-docx / Pillow 时优雅返回 []（docx 仍可走 fitz 文本路径，不崩）。
    """
    import io
    try:
        from docx import Document
        from PIL import Image
    except Exception:
        return []
    out: List[Tuple[bytes, str]] = []
    for part in Document(str(path)).part.related_parts.values():
        if not part.content_type.startswith("image/"):
            continue
        try:
            im = Image.open(io.BytesIO(part.blob))
            w, h = im.size
        except Exception:
            continue                                   # 矢量图(emf/wmf)等打不开 → 跳过
        ext = _RASTER.get((im.format or "").upper())
        if ext and w * h >= _MIN_IMG_AREA:
            out.append((part.blob, ext))
    return out


# 发票要素标记（EN+CN）：文本里没有这些、却有内嵌图 → 真发票多半在图里，应走 OCR
_INV_MARKER = re.compile(
    r"invoice\s*(no|number|#)|bill\s*no|total|amount\s+due|grand\s+total|subtotal|"
    r"发票|账单|合计|总额|总计|金额|应付|税额", re.IGNORECASE)


def _text(path) -> str:
    from docx import Document
    return "\n".join(_block_text(b) for b in _blocks(Document(str(path))))


def _text_len(path) -> int:
    return len(_text(path))


def is_image_form(path) -> bool:
    """是否"图片形式发票"：含 ≥1 张位图发票图，且 文本很少 或 文本里无发票要素（纯说明/免责声明——
    真发票其实在图里）。这样"图+大段说明文字"的 docx 也会走 OCR，不会漏掉图里的发票。缺库安全 False。"""
    try:
        if len(extract_images(path)) < 1:
            return False
        txt = _text(path)
        return len(txt) < _MIN_TEXT_CHARS or not _INV_MARKER.search(txt)
    except Exception:
        return False


def _blocks(doc):
    """按文档顺序取 body 下的段落/表格元素。"""
    from docx.oxml.ns import qn
    return [c for c in doc.element.body.iterchildren()
            if c.tag in (qn("w:p"), qn("w:tbl"))]


def _block_text(blk) -> str:
    """块内纯文本（只取 w:t 文本节点；itertext() 会重复计数、不可用）。"""
    from docx.oxml.ns import qn
    return "".join(t.text or "" for t in blk.iter(qn("w:t")))


def invoice_units(path) -> Optional[List[Tuple[int, int]]]:
    """检测 docx 内多张发票 → 各发票的块索引范围 [(b0,b1), ...]（len≥2 才拆）；否则 None。

    同 PDF 三层：TOTAL DUE 数=发票数；按 INVOICE 标题/发票号锚点定起始块（首段含起始前的前言）；
    每段须恰含 1 个 TOTAL，否则判边界不可靠、返回 None（绝不误拆）。
    """
    try:
        from ..parse.template_rules import count_total_markers
        from docx import Document
    except Exception:
        return None                                  # 缺 python-docx → 不拆，走 fitz 单张
    blocks = _blocks(Document(str(path)))
    texts = [_block_text(b) for b in blocks]
    n = count_total_markers("\n".join(texts))
    if n < 2:
        return None
    title_starts = [i for i, t in enumerate(texts) if _TITLE_RE.match(t.strip())]
    no_starts = [i for i, t in enumerate(texts) if _NO_RE.search(t)]
    starts = title_starts if len(title_starts) == n else (no_starts if len(no_starts) == n else None)
    if not starts or len(starts) < 2:
        return None
    bounds = [0] + starts[1:] + [len(blocks)]            # 起始前的前言并入首张
    ranges = [(bounds[i], bounds[i + 1] - 1) for i in range(len(bounds) - 1)]
    for a, b in ranges:
        if count_total_markers("\n".join(texts[a:b + 1])) != 1:
            return None
    return ranges


def split_docx(path, base_name: str, units: List[Tuple[int, int]]) -> List[Tuple]:
    """把多发票 docx 物理拆成一文件一发票的 docx（deepcopy 块、保留格式），落盘 uploads。

    返回 [(落盘路径, 文件名, 文件哈希), ...]；哈希由 原文件哈希+块范围 确定性派生（重处理 UPSERT 去重）。
    """
    import copy
    from docx import Document
    from docx.oxml.ns import qn
    from core import config, storage
    src = Document(str(path))
    blocks = _blocks(src)
    stem = Path(base_name).stem
    src_hash = storage.sha256_of_file(path)
    n = len(units)
    out: List[Tuple] = []
    for idx, (b0, b1) in enumerate(units, start=1):
        new = Document()
        nbody = new.element.body
        for c in list(nbody.iterchildren()):             # 清掉模板默认空段落
            if c.tag == qn("w:p"):
                nbody.remove(c)
        sect = nbody.find(qn("w:sectPr"))                # 节属性须留在末尾
        for blk in blocks[b0:b1 + 1]:
            el = copy.deepcopy(blk)
            sect.addprevious(el) if sect is not None else nbody.append(el)
        sub_hash = storage.sha256_of_bytes(f"{src_hash}:b{b0}-{b1}".encode())
        sub_name = f"{stem}_发票{idx}of{n}.docx"
        dest = config.UPLOAD_DIR / f"{sub_hash[:12]}_{sub_name}"
        new.save(dest)
        out.append((dest, sub_name, sub_hash))
    return out
